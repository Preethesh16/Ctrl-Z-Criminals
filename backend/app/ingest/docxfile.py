"""DOCX statement extraction: tables first, text-line regex fallback."""

from pathlib import Path


def read_docx_grid(path: str | Path) -> tuple[list[list], str]:
    """Returns (grid, full_text) — grid from tables, text for header meta."""
    import docx

    document = docx.Document(str(path))
    grid: list[list] = []
    for table in document.tables:
        for row in table.rows:
            grid.append([cell.text for cell in row.cells])
    text = "\n".join(p.text for p in document.paragraphs)
    return grid, text
