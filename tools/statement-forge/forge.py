"""statement-forge — synthetic fraud-case generator for TraceNet.

Generates ONE coherent investment-fraud case across 10 accounts and writes
each account's statement in a different real-world format. This is the ONLY
data allowed in public demos, committed fixtures, and screenshots
(CLAUDE.md — the police dataset is local-only).

Planted patterns (ground truth in case_manifest.json):
- Victim pays 6 UPI transfers to mule-1 (smurfing: each < Rs 50,000)
- Layering: mule-1 → mule-2 → mule-3 (rapid in-out within hours)
- ROUND TRIP: mule-3 → mule-4 → mule-5 → back to mule-1 (temporal loop)
- Cash-out: ~40% of mule-2/3 funds withdrawn at ATMs (odd hours)
- One failed/reversed IMPS transfer (identical amount credited back, REV)
- Cheque + POS noise so disposition percentages are non-trivial

Usage:  python tools/statement-forge/forge.py [outdir]
Deterministic: seeded RNG, same case every run (stable tests).
"""

import csv
import json
import random
import sys
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from decimal import Decimal
from pathlib import Path

OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).parent / "out"
RNG = random.Random(20260702)

BANKS = ["HDFC", "SBI", "AXIS", "KOTAK"]


@dataclass
class Account:
    name: str
    number: str
    bank: str
    vpa: str
    opening: Decimal
    events: list = field(default_factory=list)  # (dt, narration, debit, credit)
    rows: list = field(default_factory=list)  # (dt, narration, debit, credit, balance)

    def add(self, dt: datetime, narration: str, debit: Decimal | None, credit: Decimal | None):
        self.events.append((dt, narration, debit, credit))

    def finalize(self):
        """Sort by time FIRST, then compute the running balance — otherwise
        interleaved events would corrupt the chain."""
        self.events.sort(key=lambda e: e[0])
        balance = self.opening
        self.rows = []
        for dt, narration, debit, credit in self.events:
            balance += (credit or Decimal(0)) - (debit or Decimal(0))
            self.rows.append((dt, narration, debit, credit, balance))


def rupees(n: int) -> Decimal:
    return Decimal(n).quantize(Decimal("0.01"))


def rrn() -> str:
    return str(RNG.randint(10**11, 10**12 - 1))


def upi_pair(sender: Account, receiver: Account, dt: datetime, amount: Decimal, note: str = "payment"):
    ref = rrn()
    sender.add(dt, f"UPI/DR/{ref}/{receiver.name.upper()}/{receiver.bank[:4]}/{receiver.vpa}/{note}", amount, None)
    receiver.add(dt + timedelta(minutes=RNG.randint(0, 3)),
                 f"UPI/CR/{ref}/{sender.name.upper()}/{sender.bank[:4]}/{sender.vpa}/{note}", None, amount)
    return ref


def imps_pair(sender: Account, receiver: Account, dt: datetime, amount: Decimal):
    ref = rrn()
    sender.add(dt, f"IMPS/P2A/{ref}/{receiver.name.upper()}/{receiver.number[-4:]}", amount, None)
    receiver.add(dt + timedelta(minutes=RNG.randint(1, 20)),
                 f"IMPS/P2A/{ref}/{sender.name.upper()}/CR", None, amount)
    return ref


def build_case():
    victim = Account("Suresh Patil", "50100234567891", "HDFC", "suresh.p@okhdfc", rupees(2_400_000))
    mules = [
        Account(f"Mule{i} {n}", f"3{i}00{RNG.randint(10**9, 10**10-1)}", BANKS[i % 4], f"{n.lower()}{i}@ok{BANKS[i % 4].lower()}", rupees(RNG.randint(2_000, 20_000)))
        for i, n in enumerate(
            ["Ramesh", "Vikram", "Anita", "Farid", "Deepak", "Sonal", "Kiran", "Manoj"], start=1)
    ]
    m = {i + 1: a for i, a in enumerate(mules)}
    manifest: dict = {"accounts": {}, "planted": {}}

    t0 = datetime(2026, 5, 4, 10, 15)

    # --- background noise on victim account
    for d in range(1, 25, 3):
        victim.add(t0 - timedelta(days=30 - d), f"POS 411111XXXXXX1111 BIGBASKET BLR {rrn()}", rupees(RNG.randint(500, 4000)), None)
    victim.add(t0 - timedelta(days=20), f"NEFT-HDFCN5{RNG.randint(10**13, 10**14-1)}-ACME SALARY", None, rupees(180_000))

    # --- 1. smurfing: victim → mule1, six UPI txns each < 50k, same day
    smurf_refs = []
    for k in range(6):
        amt = rupees(RNG.randint(41_000, 49_500))
        smurf_refs.append(upi_pair(victim, m[1], t0 + timedelta(minutes=25 * k), amt, "invest plan"))
    manifest["planted"]["smurfing"] = {"from": victim.number, "to": m[1].number, "refs": smurf_refs}

    # --- 2. layering: mule1 → mule2 → mule3 (rapid in-out)
    l1 = imps_pair(m[1], m[2], t0 + timedelta(hours=3), rupees(150_000))
    l2 = imps_pair(m[1], m[2], t0 + timedelta(hours=3, minutes=40), rupees(90_000))
    l3 = imps_pair(m[2], m[3], t0 + timedelta(hours=6), rupees(140_000))
    manifest["planted"]["layering"] = [l1, l2, l3]

    # --- 3. ROUND TRIP: mule3 → mule4 → mule5 → mule1 (time-ordered loop)
    r1 = upi_pair(m[3], m[4], t0 + timedelta(days=1, hours=2), rupees(60_000), "goods")
    r2 = upi_pair(m[4], m[5], t0 + timedelta(days=1, hours=5), rupees(58_000), "repay")
    r3 = upi_pair(m[5], m[1], t0 + timedelta(days=1, hours=9), rupees(56_500), "gift")
    manifest["planted"]["round_trip"] = {
        "path": [m[3].number, m[4].number, m[5].number, m[1].number],
        "refs": [r1, r2, r3],
        "note": "m1 received smurfed funds; loop returns to m1 — linked cluster closes",
    }

    # --- 4. cash-out ~40% at ATMs, odd hours
    atm_total = Decimal(0)
    for k, (who, amt) in enumerate([(2, 40_000), (2, 20_000), (3, 30_000), (3, 20_000)]):
        dt = t0 + timedelta(days=1 + k, hours=1, minutes=RNG.randint(0, 50))  # 01:xx
        m[who].add(dt, f"ATM WDL {dt:%H:%M} MG ROAD BLR {rrn()}", rupees(amt), None)
        atm_total += rupees(amt)
    manifest["planted"]["cash_out_total"] = str(atm_total)

    # --- 5. one failed transfer, reversed next day
    ref = imps_pair(m[2], m[6], t0 + timedelta(days=2), rupees(25_000))
    m[2].add(t0 + timedelta(days=3), f"IMPS REV {ref} RETURNED FAILED CREDIT", None, rupees(25_000))
    manifest["planted"]["reversal"] = {"ref": ref, "amount": "25000.00"}

    # --- 6. cheque + noise
    m[1].add(t0 + timedelta(days=2, hours=4), f"CHQ NO {RNG.randint(100000, 999999)} CLEARING SELF", rupees(35_000), None)
    for i in (4, 5, 6, 7, 8):
        m[i].add(t0 + timedelta(days=RNG.randint(3, 6)), f"POS {rrn()} DMART", rupees(RNG.randint(300, 2500)), None)

    accounts = [victim, *mules]
    for a in accounts:
        a.finalize()
        manifest["accounts"][a.number] = {"name": a.name, "bank": a.bank, "vpa": a.vpa,
                                          "opening": str(a.opening), "rows": len(a.rows)}
    return accounts, manifest


# ---------------- writers: each account → a different real-world format ----

def write_csv(a: Account, path: Path):
    with path.open("w", newline="") as f:
        w = csv.writer(f)
        w.writerow([f"{a.bank} BANK — STATEMENT", "", "", "", "", ""])
        w.writerow([f"Account No: {a.number}", "", f"Name: {a.name}", "", "", ""])
        w.writerow(["TRAN-DATE", "TRAN_PARTICULAR", "CHQ-NUM", "WITHDRAWAL", "DEPOSIT", "BALANCE"])
        for dt, narr, dr, cr, bal in a.rows:
            w.writerow([dt.strftime("%d-%m-%Y"), narr, "", dr or "", cr or "", bal])


def write_xlsx(a: Account, path: Path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append([f"{a.bank} BANK"])
    ws.append([f"Account No: {a.number}  Customer Name: {a.name}"])
    ws.append([])
    ws.append(["Tran Date", "Particulars", "Debit", "Credit", "Balance"])
    for dt, narr, dr, cr, bal in a.rows:
        ws.append([dt.strftime("%d/%m/%Y %H:%M"), narr,
                   float(dr) if dr else None, float(cr) if cr else None, float(bal)])
    wb.save(path)


def write_pdf(a: Account, path: Path):
    """Digital PDF with a ruled table (pdfplumber-extractable)."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4, title=f"{a.bank} statement")
    head = Paragraph(
        f"<b>{a.bank} BANK</b><br/>Account No: {a.number}<br/>Customer Name: {a.name}<br/>"
        f"Statement Period: {a.rows[0][0]:%d-%m-%Y} to {a.rows[-1][0]:%d-%m-%Y}",
        styles["Normal"],
    )
    body = [["Date", "Narration", "Debit", "Credit", "Balance"]]
    for dt, narr, dr, cr, bal in a.rows:
        body.append([dt.strftime("%d-%m-%Y"), Paragraph(narr, styles["BodyText"]),
                     f"{dr:.2f}" if dr else "", f"{cr:.2f}" if cr else "", f"{bal:.2f}"])
    table = Table(body, colWidths=[22 * mm, 88 * mm, 22 * mm, 22 * mm, 24 * mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
    ]))
    doc.build([head, Spacer(1, 6 * mm), table])


def write_txt(a: Account, path: Path):
    lines = [f"{a.bank} BANK", f"Account No : {a.number}", f"Name : {a.name}", ""]
    lines.append(f"{'Trans Dt':<12}  {'Transaction Particulars':<58}  {'Debit':>12}  {'Credit':>12}  {'Balance':>14}")
    for dt, narr, dr, cr, bal in a.rows:
        lines.append(
            f"{dt.strftime('%d-%m-%Y'):<12}  {narr[:58]:<58}  "
            f"{(f'{dr:.2f}' if dr else ''):>12}  {(f'{cr:.2f}' if cr else ''):>12}  {bal:>14.2f}"
        )
    path.write_text("\n".join(lines))


def write_html_xls(a: Account, path: Path):
    """Bank-style '.xls' that is actually an HTML table."""
    rows = "".join(
        f"<tr><td>{dt:%d-%m-%Y}</td><td>{narr}</td><td>{dr or ''}</td><td>{cr or ''}</td><td>{bal}</td></tr>"
        for dt, narr, dr, cr, bal in a.rows
    )
    path.write_text(
        f"<html><body><h3>{a.bank} BANK — {a.number} — {a.name}</h3>"
        f"<table border=1><tr><th>Tran Date</th><th>Particulars</th><th>Debit</th>"
        f"<th>Credit</th><th>Balance</th></tr>{rows}</table></body></html>"
    )


def write_docx(a: Account, path: Path):
    import docx

    d = docx.Document()
    d.add_heading(f"{a.bank} BANK", level=2)
    d.add_paragraph(f"Account No: {a.number}")
    d.add_paragraph(f"Customer Name: {a.name}")
    table = d.add_table(rows=1, cols=5)
    for i, h in enumerate(["Tran Date", "Particulars", "Debit", "Credit", "Balance"]):
        table.rows[0].cells[i].text = h
    for dt, narr, dr, cr, bal in a.rows:
        cells = table.add_row().cells
        cells[0].text = dt.strftime("%d-%m-%Y")
        cells[1].text = narr
        cells[2].text = f"{dr:.2f}" if dr else ""
        cells[3].text = f"{cr:.2f}" if cr else ""
        cells[4].text = f"{bal:.2f}"
    d.save(str(path))


def write_scanned_pdf(a: Account, path: Path):
    """Image-only PDF (no text layer) — exercises the OCR pipeline.
    Rendered from the digital PDF at 200 DPI, like a branch scan."""
    import tempfile

    from pdf2image import convert_from_path

    with tempfile.NamedTemporaryFile(suffix=".pdf") as tmp:
        write_pdf(a, Path(tmp.name))
        pages = convert_from_path(tmp.name, dpi=200)
    # resolution must match the render DPI or the PDF page becomes 23x32in
    pages[0].save(str(path), save_all=True, append_images=pages[1:], resolution=200)


WRITERS = [
    ("victim_hdfc.pdf", write_pdf),
    ("mule1_sbi.csv", write_csv),
    ("mule2_axis.xlsx", write_xlsx),
    ("mule3_kotak.pdf", write_pdf),
    ("mule4_hdfc.xls", write_html_xls),
    ("mule5_sbi.txt", write_txt),
    ("mule6_axis.docx", write_docx),
    ("mule7_kotak.xlsx", write_xlsx),
    ("mule8_hdfc_scanned.pdf", write_scanned_pdf),
]


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for stale in OUT.glob("*"):  # deterministic output — remove previous run
        if stale.suffix in (".pdf", ".csv", ".xlsx", ".xls", ".txt", ".docx", ".json"):
            stale.unlink()
    accounts, manifest = build_case()
    for account, (fname, writer) in zip(accounts, WRITERS):
        writer(account, OUT / fname)
        manifest["accounts"][account.number]["file"] = fname
    (OUT / "case_manifest.json").write_text(json.dumps(manifest, indent=1, default=str))
    total = sum(len(a.rows) for a in accounts)
    print(f"wrote {len(accounts)} statements ({total} txns) + case_manifest.json → {OUT}")


if __name__ == "__main__":
    main()
